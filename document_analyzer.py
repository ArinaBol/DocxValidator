from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm

class DocumentAnalyzer:
    def __init__(self, file_path):
        self.doc = Document(file_path)
        self.report = []

    def check_text_format(self, expected_font, expected_size, indent_min, indent_max, alignment, spacing_before, spacing_after):
        for i, paragraph in enumerate(tqdm(self.doc.paragraphs, desc="Перевірка форматування тексту"), 1):
            for run in paragraph.runs:
                font = run.font
                size = font.size.pt if font.size else None

                if font.name != expected_font or size != expected_size:
                    self.report.append({
                        "type": "Невідповідність шрифту",
                        "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                        "actual_value": f"{font.name}, {size}pt" if font.name else "Невідомий шрифт",
                        "expected_value": f"{expected_font}, {expected_size}pt",
                        "location": f"Абзац {i}"
                    })

            if paragraph.alignment != alignment:
                self.report.append({
                    "type": "Невідповідність вирівнювання",
                    "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                    "actual_value": str(paragraph.alignment),
                    "expected_value": "По ширині",
                    "location": f"Абзац {i}"
                })

            indent = paragraph.paragraph_format.first_line_indent
            indent_value = round(indent.cm, 2) if indent else 0
            if not (indent_min <= indent_value <= indent_max):
                self.report.append({
                    "type": "Невідповідність абзацного відступу",
                    "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                    "actual_value": f"{indent_value} см",
                    "expected_value": f"{indent_min}-{indent_max} см",
                    "location": f"Абзац {i}"
                })

            space_before = paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0
            space_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0
            if space_before != spacing_before or space_after != spacing_after:
                self.report.append({
                    "type": "Невідповідність міжабзацних відступів",
                    "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                    "actual_value": f"До: {space_before}, Після: {space_after}",
                    "expected_value": f"До: {spacing_before}, Після: {spacing_after}",
                    "location": f"Абзац {i}"
                })

    def check_table_of_contents(self):
        headings = [p.text for p in self.doc.paragraphs if p.style.name.startswith("Heading")]
        for i, heading in enumerate(tqdm(headings, desc="Перевірка змісту"), 1):
            if f"{i}." not in heading:
                self.report.append({
                    "type": "Невідповідність змісту",
                    "fragment": heading[:50] + "...",
                    "actual_value": "Відсутній або неправильний номер",
                    "expected_value": f"Розділ {i}",
                    "location": "Зміст"
                })

    def check_references(self):
        references_section = False
        for i, paragraph in enumerate(tqdm(self.doc.paragraphs, desc="Перевірка списку джерел"), 1):
            if "Список джерел" in paragraph.text:
                references_section = True
            if references_section and not paragraph.text.strip().isdigit():
                self.report.append({
                    "type": "Невідповідність у списку джерел",
                    "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                    "actual_value": "Неправильний формат посилання",
                    "expected_value": "Послідовна нумерація",
                    "location": f"Абзац {i}"
                })

    def check_figures(self):
        for i, paragraph in enumerate(tqdm(self.doc.paragraphs, desc="Перевірка рисунків"), 1):
            if "Рисунок" in paragraph.text:
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.report.append({
                        "type": "Невідповідність вирівнювання рисунка",
                        "fragment": " ".join(paragraph.text.split()[:5]) + "...",
                        "actual_value": "Не по центру",
                        "expected_value": "По центру",
                        "location": f"Абзац {i}"
                    })

    def generate_report(self, output_file):
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("Звіт про аналіз документу
")
            f.write("="*50 + "

")
            for entry in self.report:
                f.write(f"Тип помилки: {entry['type']}
")
                f.write(f"Фрагмент: {entry['fragment']}
")
                f.write(f"Фактичне значення: {entry['actual_value']}
")
                f.write(f"Очікуване значення: {entry['expected_value']}
")
                f.write(f"Розташування: {entry['location']}
")
                f.write("-"*50 + "
")

# Використання
file_path = "/content/Курсова_МКО_Болотнікова_О.Д..docx"  # Назва завантаженого файлу
output_file = "analysis_report.txt"

analyzer = DocumentAnalyzer(file_path)
analyzer.check_text_format("Times New Roman", 14, 1.25, 1.5, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 0)
analyzer.check_table_of_contents()
analyzer.check_references()
analyzer.check_figures()
analyzer.generate_report(output_file)

print(f"Звіт збережено у {output_file}")
    
    